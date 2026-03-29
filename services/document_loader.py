"""Document loaders — extract text from various file types."""
import os
from typing import List

from langchain_core.documents import Document


def load_pdf(file_path: str) -> List[Document]:
    from langchain_community.document_loaders import PyPDFLoader
    loader = PyPDFLoader(file_path)
    return loader.load()


def load_docx(file_path: str) -> List[Document]:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return [Document(page_content=text, metadata={"source": file_path})]


def load_txt(file_path: str) -> List[Document]:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="cp1254") as f:
            text = f.read()
    return [Document(page_content=text, metadata={"source": file_path})]


def load_excel(file_path: str) -> List[Document]:
    import pandas as pd
    dfs = pd.read_excel(file_path, sheet_name=None)
    documents = []
    for sheet_name, df in dfs.items():
        text = f"Sayfa: {sheet_name}\n{df.to_string(index=False)}"
        documents.append(Document(page_content=text, metadata={"source": file_path, "sheet": sheet_name}))
    return documents


def load_document(file_path: str, file_type: str) -> List[Document]:
    """Load a document based on file type and return LangChain Document objects."""
    loaders = {
        "pdf": load_pdf,
        "docx": load_docx,
        "doc": load_docx,
        "txt": load_txt,
        "xlsx": load_excel,
        "xls": load_excel,
        "csv": load_txt,
    }

    loader_fn = loaders.get(file_type.lower())
    if not loader_fn:
        raise ValueError(f"Desteklenmeyen dosya türü: {file_type}")

    docs = loader_fn(file_path)
    # Add file_name metadata
    for doc in docs:
        doc.metadata["file_name"] = os.path.basename(file_path)
    return docs
